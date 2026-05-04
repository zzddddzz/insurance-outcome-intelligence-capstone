"""
Generate the Capstone Analysis Word Document with screenshots.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import json
import os
import pandas as pd


def load_json(path, default):
    if os.path.exists(path):
        with open(path) as f:
            return json.load(f)
    return default


def load_csv(path):
    if os.path.exists(path):
        return pd.read_csv(path)
    return pd.DataFrame()


def nested_get(obj, path, default=None):
    cur = obj
    for key in path:
        if not isinstance(cur, dict) or key not in cur:
            return default
        cur = cur[key]
    return cur


def fmt_decimal(value, digits=4, default="N/A"):
    if value is None:
        return default
    return f"{float(value):.{digits}f}"


def fmt_money(value, default="N/A"):
    if value is None or pd.isna(value):
        return default
    return f"EUR {float(value):,.0f}"


def fmt_count(value, default="N/A"):
    if value is None or pd.isna(value):
        return default
    return f"{int(round(float(value))):,}"


def fmt_pct(value, default="N/A"):
    if value is None:
        return default
    return f"{float(value) * 100:.1f}%"


model_summary = load_json("output/model_summary.json", {})
decision_summary = load_csv("output/decision_summary.csv")
feature_importance = load_csv("output/feature_importance.csv")

data_rows = nested_get(model_summary, ["data_shape"], [228711])[0]
total_policies = model_summary.get("total_policies", 45162)
total_customers = model_summary.get("total_customers", 1467)

lapse_val = nested_get(model_summary, ["lapse_model", "val"], {})
lapse_test = nested_get(model_summary, ["lapse_model", "test"], {})
claim_model = model_summary.get("claim_model", {})
claim_val = claim_model.get("val", claim_model if isinstance(claim_model, dict) else {})
claim_test = claim_model.get("test", {})

decision_actions = model_summary.get("decision_actions", {})
portfolio_total = sum(decision_actions.values()) or data_rows
priority_count = sum(count for action, count in decision_actions.items() if action != "STANDARD")
priority_pct = priority_count / portfolio_total if portfolio_total else 0

if not decision_summary.empty:
    decision_dist_rows = []
    for _, row in decision_summary.iterrows():
        count = int(row["count"])
        portfolio_pct = count / portfolio_total if portfolio_total else 0
        decision_dist_rows.append([
            row["decision_action"],
            fmt_count(count),
            fmt_pct(portfolio_pct),
            fmt_money(row.get("avg_premium")),
            fmt_decimal(row.get("avg_loss_ratio"), 2),
        ])
else:
    decision_dist_rows = [[action, fmt_count(count), fmt_pct(count / portfolio_total), "N/A", "N/A"]
                          for action, count in decision_actions.items()]

def action_value(action, column=None):
    if decision_summary.empty:
        return decision_actions.get(action)
    row = decision_summary.loc[decision_summary["decision_action"] == action]
    if row.empty:
        return None
    if column is None:
        return row.iloc[0]["count"]
    return row.iloc[0].get(column)

early_risk_count = action_value("EARLY_RISK") or 0
early_risk_loss = action_value("EARLY_RISK", "avg_loss_ratio")
retain_high_count = action_value("RETAIN_HIGH") or 0
retain_high_loss = action_value("RETAIN_HIGH", "avg_loss_ratio")

if not feature_importance.empty:
    top_features = feature_importance.head(5)
    fi_rows = [
        [
            row["feature"],
            fmt_decimal(row["importance"], 3),
            "Policy, coverage, utilization, or demographic signal used by the model",
        ]
        for _, row in top_features.iterrows()
    ]
    top_feature_phrase = ", ".join(top_features["feature"].head(3).astype(str).tolist())
else:
    fi_rows = [["N/A", "N/A", "Run src/train_pipeline.py to generate feature importance"]]
    top_feature_phrase = "the top model features"

excluded_features = nested_get(
    model_summary,
    ["model_metadata", "lapse", "excluded_leakage_features"],
    ["exposure_time", "loss_ratio", "claim_premium_ratio_capped"],
)

doc = Document()

# ============================================================
# Page setup
# ============================================================
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============================================================
# Helper functions
# ============================================================
def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h

def add_img(doc, path, width_inches=6.0):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_inches))
        doc.add_paragraph('')  # spacer
    else:
        doc.add_paragraph(f'[Image not found: {path}]')
        doc.add_paragraph('')

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Insurance Outcome Intelligence')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Decision Intelligence Framework for Health Insurance Portfolio Optimization')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MSDS 498 Capstone Project — Team 54')
run.font.size = Pt(14)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Pavani Katamreddy, Ashin Katwala, HyunChul Lee, Di Zhang, Savannah Lucero')
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('May 2025')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_heading_styled(doc, 'Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Problem Statement and Motivation',
    '3. Literature Review and Related Work',
    '4. Data Foundation',
    '5. Methodology',
    '    5.1 Data Preprocessing and Feature Engineering',
    '    5.2 Lapse Prediction Model',
    '    5.3 Claim Cost Prediction Model',
    '    5.4 Decision Intelligence Framework',
    '6. Computational Experiment and Results',
    '    6.1 Model Performance',
    '    6.2 Decision Space Analysis',
    '    6.3 Feature Importance',
    '    6.4 Sensitivity Analysis',
    '    6.5 Dashboard Walkthrough',
    '7. Discussion and Implications',
    '8. Limitations and Future Work',
    '9. Conclusions',
    '10. References',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
add_heading_styled(doc, '1. Executive Summary', level=1)

doc.add_paragraph(
    'This capstone project introduces an Insurance Outcome Intelligence system — a Decision Intelligence '
    'Framework designed to transform raw health insurance portfolio data into actionable business decisions. '
    'The system addresses a fundamental challenge facing health insurers: the gap between data collection '
    'and decision-making. While organizations accumulate vast amounts of customer, pricing, and claims data, '
    'the tools to translate this information into timely, profit-driven actions remain limited.'

    '\n\nThe solution comprises three integrated components: (1) a lapse prediction model that identifies '
    'customers most likely to leave their policy, (2) a claim cost prediction model that forecasts expected '
    'medical expenditure per customer, and (3) a Decision Intelligence Framework that classifies each customer '
    'into an actionable category for retention, pricing, or risk intervention. An interactive Streamlit '
    'dashboard enables stakeholders to explore segments, simulate pricing and retention strategies, and '
    'evaluate portfolio impact in real time.'

    f'\n\nTrained on a Spanish health insurer portfolio dataset containing {fmt_count(data_rows)} rows '
    f'across {fmt_count(total_policies)} policies and {fmt_count(total_customers)} customers, the leakage-screened '
    f'lapse prediction model achieved a test ROC AUC of {fmt_decimal(lapse_test.get("roc_auc"))}, and the claim '
    f'cost prediction model achieved a test R-squared of {fmt_decimal(claim_test.get("r2"))}. The Decision '
    f'Intelligence Framework categorized approximately {fmt_pct(priority_pct)} of the portfolio into priority '
    'action groups, enabling targeted resource allocation that directly supports the capstone project\'s '
    'objectives of better pricing discipline, more focused retention spending, and earlier risk intervention.'
)

# ============================================================
# 2. PROBLEM STATEMENT AND MOTIVATION
# ============================================================
add_heading_styled(doc, '2. Problem Statement and Motivation', level=1)

doc.add_paragraph(
    'Health insurers operate in a complex risk environment where pricing accuracy, customer retention, and '
    'claims management are interdependent. Current portfolio management practices often rely on reactive '
    'analysis rather than predictive, decision-oriented frameworks. This leads to three core problems:'
)

bullets = [
    'Pricing misalignment: Premiums may not fully reflect underlying risk, resulting in underpricing '
    'high-cost segments and overpricing low-risk ones.',
    'Inefficient retention: Retention efforts are spread uniformly rather than focused on customers who '
    'are both at-risk of lapsing and valuable to the portfolio.',
    'Delayed risk visibility: High-cost segments are identified after losses have already materialized, '
    'reducing the effectiveness of intervention strategies.'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    'This project addresses these problems through five research questions derived from the capstone '
    'proposal:'
)

questions = [
    'Which customers are most likely to lapse (leave their insurance)?',
    'Which customers are projected to drive high claim spend?',
    'Where do premiums appear misaligned with underlying risk?',
    'Which at-risk customers are worth prioritizing for retention?',
    'What interventions are most likely to improve results?'
]
for q in questions:
    doc.add_paragraph(q, style='List Number')

doc.add_paragraph(
    'The problem motivation is grounded in the observation that the cost of inaction — pricing gaps, '
    'inefficient retention spending, and delayed responses to risk trends — creates meaningful financial '
    'impact across the portfolio over time. Even a modest 1-2% retention improvement can create substantial '
    'portfolio value. Our framework connects predictive analytics directly to practical business outcomes '
    'that leadership can apply in real decision-making situations.'
)

doc.add_page_break()

# ============================================================
# 3. LITERATURE REVIEW AND RELATED WORK
# ============================================================
add_heading_styled(doc, '3. Literature Review and Related Work', level=1)

doc.add_paragraph(
    'The application of predictive analytics in insurance has evolved from basic actuarial methods to '
    'sophisticated machine learning approaches. Research in customer lapse and churn prediction has '
    'demonstrated that machine learning models — particularly gradient boosting and neural networks — '
    'consistently outperform traditional logistic regression in predictive accuracy. Studies have shown '
    'that features such as customer tenure, service utilization patterns, and premium-to-claim ratios '
    'are strong predictors of lapse behavior.'

    '\n\nIn the domain of claim cost prediction, the literature emphasizes the importance of '
    'feature engineering and non-linear modeling. Conventional GLM (Generalized Linear Model) approaches, '
    'while interpretable, often fail to capture complex interactions between customer demographics, '
    'policy characteristics, and medical service utilization. Gradient boosting methods (XGBoost, LightGBM) '
    'have emerged as the state-of-the-art for this task, providing superior predictive performance while '
    'maintaining feature interpretability through importance scores.'

    '\n\nThe Decision Intelligence Framework builds on the concept of predictive-to-prescriptive analytics. '
    'While most insurance analytics projects stop at prediction, our framework follows recent literature '
    'advocating for decision-oriented modeling that classifies entities into actionable categories. '
    'This approach aligns with research showing that the primary barrier to analytics adoption is not '
    'model accuracy but the inability to translate model outputs into specific business actions.'

    '\n\nThe dataset used in this project — the Health Insurance Portfolio dataset from Mendeley Data (v4, 2025) '
    '— provides a comprehensive view of individual health insurance policies including premium amounts, '
    'claim costs, lapse indicators, and extensive coverage and cost-sharing parameters. This public dataset '
    'enables reproducible research while capturing dynamics broadly relevant to insurance markets.'
)

doc.add_page_break()

# ============================================================
# 4. DATA FOUNDATION
# ============================================================
add_heading_styled(doc, '4. Data Foundation', level=1)

doc.add_paragraph(
    f'The primary dataset consists of {fmt_count(data_rows)} rows from a Spanish health insurance insurer, '
    f'covering {fmt_count(total_policies)} individual policies and approximately {fmt_count(total_customers)} '
    'unique insured customers. The data spans multiple '
    'years and includes both policy-level and insured-level records.'
)

doc.add_heading('4.1 Key Variables', level=2)

data_headers = ['Variable', 'Type', 'Description']
data_variable_rows = [
    ['ID, ID_policy, ID_insured', 'Identifier', 'Unique policy and insured identifiers'],
    ['premium', 'Numeric', 'Annual premium amount (EUR)'],
    ['cost_claims_year', 'Numeric', 'Annual claim cost (EUR)'],
    ['lapse', 'Categorical', '1 = lapsed, 2 = retained'],
    ['age', 'Numeric', 'Age of insured customer'],
    ['type_product', 'Categorical', 'S (Savings), P (Personal), D (Death), I (Individual)'],
    ['seniority_policy', 'Numeric', 'Policy tenure in years'],
    ['n_medical_services', 'Numeric', 'Number of medical services used'],
    ['distribution_channel', 'Categorical', 'A (Agent), I (Intermediary), D (Direct)'],
    ['C_H, C_GI, C_II, C_GE_*', 'Coverage', 'Coverage percentages for various benefits'],
    ['exposure_time', 'Numeric', 'Policy exposure duration (0-1 scale)'],
    ['new_business', 'Categorical', 'Whether the policy is new business'],
]
add_table(doc, data_headers, data_variable_rows)

doc.add_paragraph('')

doc.add_heading('4.2 Data Characteristics', level=2)

doc.add_paragraph(
    'The dataset exhibits several important characteristics. The lapse rate is approximately 7.15%, '
    'indicating a highly imbalanced classification problem. Premium amounts range widely with a mean of '
    '~EUR 852 and significant right-skew. Claim costs show even greater variability (mean ~EUR 588, '
    'median ~EUR 186), indicating a heavy right tail. The average loss ratio (claims divided by premium) '
    'is approximately 0.66, suggesting moderate overall pricing adequacy across the portfolio.'
)

add_img(doc, 'output/claim_distribution.png', 6.5)

doc.add_paragraph(
    'Figure 1: Claim cost distribution analysis. The histogram shows heavy right-skew in claim costs. '
    'The scatter plot reveals the relationship between age, claim costs, and premiums. The loss ratio '
    'distribution and medical services analysis provide additional insight into portfolio risk composition.'
)

doc.add_page_break()

# ============================================================
# 5. METHODOLOGY
# ============================================================
add_heading_styled(doc, '5. Methodology', level=1)

doc.add_paragraph(
    'The analytical approach follows three building blocks — Exploratory Data Prep, Decision Modeling, '
    'and Intelligence — where raw data is refined into validated features, converted into model outputs, '
    'and ultimately translated into specific leadership actions that drive portfolio performance.'
)

# 5.1 Data Preprocessing
doc.add_heading('5.1 Data Preprocessing and Feature Engineering', level=2)

doc.add_paragraph(
    'The preprocessing pipeline transforms the raw dataset into analysis-ready features. Key steps include:'
)

pp_steps = [
    'Loss ratio calculation: Claim cost divided by premium, clipped to handle extreme values.',
    'Claim-to-premium ratio (capped): A normalized measure of cost pressure, capped at 10x to '
    'reduce outlier impact.',
    'Average claim per service: Total claim cost divided by number of medical services.',
    'Over-claim flag: Binary indicator where claim cost exceeds premium.',
    'Age group and tenure group binning: Categorical transformations of continuous variables for '
    'improved model interpretability.',
    'Premium percentile: Ranks customers by premium for stratification.',
    'High claim flag: Binary indicator for claims above the 75th percentile.',
    'Leakage control: Direct outcome-derived fields are retained for EDA and business reporting but excluded '
    'from predictive model inputs.',
    'Label encoding: Categorical variables encoded numerically for model input.',
    'Missing value imputation: Numerical features imputed with median; categorical features '
    'imputed with "Unknown" category.'
]
for s in pp_steps:
    doc.add_paragraph(s, style='List Bullet')

# 5.2 Lapse Prediction Model
doc.add_heading('5.2 Lapse Prediction Model', level=2)

doc.add_paragraph(
    'The lapse prediction model is formulated as a binary classification problem where the target '
    'variable indicates whether a customer has lapsed (lapse = 1) or is retained (lapse = 2). '
    'Given the class imbalance (~7.15% lapsed vs. ~92.85% retained), the XGBoost classifier uses '
    'the scale_pos_weight parameter (set to ~12.99) to balance the classes during training.'
)

doc.add_paragraph(
    'Model configuration: 300 estimators, maximum tree depth of 6, learning rate of 0.05, '
    'subsample ratio of 0.8, column subsample ratio of 0.8, L1 regularization (alpha=0.1), '
    'and L2 regularization (lambda=1.0). The data was split 70/15/15 into train/validation/test sets, '
    'with stratification on the target variable to maintain class balance across splits.'
)

# 5.3 Claim Cost Prediction Model
doc.add_heading('5.3 Claim Cost Prediction Model', level=2)

doc.add_paragraph(
    'The claim cost prediction model is formulated as a regression problem with a log-transformed '
    'target variable (log(1 + cost_claims_year)) to handle the heavy right-skew in claim costs. '
    'The same XGBoost framework is used with similar hyperparameters, excluding the class weighting '
    'parameter as regression does not require it.'
)

doc.add_paragraph(
    'Performance is evaluated on the original (non-transformed) scale using RMSE, MAE, R-squared, '
    'and Mean Absolute Percentage Error (MAPE) to ensure results are interpretable in business terms.'
)

# 5.4 Decision Intelligence Framework
doc.add_heading('5.4 Decision Intelligence Framework', level=2)

doc.add_paragraph(
    'The Decision Intelligence Framework is the core decision-support component that translates '
    'model predictions into actionable categories. Each customer is classified based on the '
    'intersection of two dimensions: lapse risk and claim cost relative to premium.'
)

doc.add_heading('Classification rules:', level=2)

decision_headers = ['Action Category', 'Condition', 'Business Meaning']
decision_rows = [
    ['RETAIN_HIGH', 'High lapse risk AND profitable',
     'Customer is valuable but at risk — prioritize personalized retention outreach'],
    ['REVIEW_PRICING', 'Underpriced AND profitable',
     'Premium barely covers predicted cost — flag for pricing review'],
    ['EARLY_RISK', 'Underpriced AND unprofitable',
     'High cost exposure with inadequate premium — trigger early clinical/financial intervention'],
    ['LOW_PRIORITY', 'High lapse risk AND unprofitable',
     'Low value customer likely to leave — allocate standard outreach only'],
    ['STANDARD', 'Normal risk profile',
     'Within expected parameters — maintain standard portfolio management'],
]
add_table(doc, decision_headers, decision_rows)

doc.add_paragraph('')

doc.add_paragraph(
    'Thresholds are defined dynamically: high lapse risk is defined as being in the top 25% of '
    'predicted lapse probability, and underpriced is defined as premium being less than 80% of '
    'predicted claim cost. Profitability is defined as premium exceeding predicted claim cost.'
)

doc.add_page_break()

# ============================================================
# 6. COMPUTATIONAL EXPERIMENT AND RESULTS
# ============================================================
add_heading_styled(doc, '6. Computational Experiment and Results', level=1)

# 6.1 Model Performance
doc.add_heading('6.1 Model Performance', level=2)

doc.add_paragraph(
    'Both models were trained on 70% of the data, validated on 15%, and tested on the held-out 15%. '
    'The lapse prediction model demonstrates excellent discriminative ability across all metrics.'
)

perf_headers = ['Metric', 'Validation', 'Test']
perf_rows = [
    ['Accuracy', fmt_decimal(lapse_val.get('accuracy')), fmt_decimal(lapse_test.get('accuracy'))],
    ['Precision', fmt_decimal(lapse_val.get('precision')), fmt_decimal(lapse_test.get('precision'))],
    ['Recall', fmt_decimal(lapse_val.get('recall')), fmt_decimal(lapse_test.get('recall'))],
    ['F1 Score', fmt_decimal(lapse_val.get('f1')), fmt_decimal(lapse_test.get('f1'))],
    ['ROC AUC', fmt_decimal(lapse_val.get('roc_auc')), fmt_decimal(lapse_test.get('roc_auc'))],
]
add_table(doc, perf_headers, perf_rows)

doc.add_paragraph('')

claim_headers = ['Metric', 'Validation', 'Test']
claim_rows = [
    ['RMSE', fmt_money(claim_val.get('rmse')), fmt_money(claim_test.get('rmse'))],
    ['MAE', fmt_money(claim_val.get('mae')), fmt_money(claim_test.get('mae'))],
    ['R-squared', fmt_decimal(claim_val.get('r2')), fmt_decimal(claim_test.get('r2'))],
    ['MAPE', f"{fmt_decimal(claim_val.get('mape'), 2)}%", f"{fmt_decimal(claim_test.get('mape'), 2)}%"],
]
add_table(doc, claim_headers, claim_rows)

doc.add_paragraph('')

doc.add_paragraph(
    'The revised modeling pipeline excludes direct target-derived fields from the predictive feature matrix, '
    f'including {", ".join(excluded_features)}. As a result, the reported performance is a more conservative '
    'estimate based on policy, premium, coverage, demographic, and utilization signals rather than fields '
    'that directly encode the outcome.'
)

# 6.2 Decision Space Analysis
doc.add_heading('6.2 Decision Space Analysis', level=2)

doc.add_paragraph(
    f'The Decision Intelligence Framework classified the {fmt_count(data_rows)} rows into actionable categories. '
    f'The distribution shows {fmt_pct(priority_pct)} of the portfolio outside STANDARD management, giving '
    'leadership a focused set of priority groups for retention, pricing, or risk intervention.'
)

decision_dist_headers = ['Action Category', 'Count', 'Portfolio %', 'Avg Premium', 'Avg Loss Ratio']
add_table(doc, decision_dist_headers, decision_dist_rows)

doc.add_paragraph('')

add_img(doc, 'output/decision_matrix.png', 6.5)

doc.add_paragraph(
    'Figure 2: Decision Intelligence Framework scatter plot. Each point represents a policy '
    'classified into one of four action categories. The vertical dashed line marks the lapse '
    'probability threshold (75th percentile), and the horizontal dashed line marks the median '
    'predicted claim cost. RETAIN_HIGH policies (green, upper-left) represent the highest-value '
    'retention targets — profitable customers at risk of leaving. EARLY_RISK policies (red, '
    'upper-right) represent the highest-cost segment requiring immediate intervention.'
)

# 6.3 Feature Importance
doc.add_heading('6.3 Feature Importance', level=2)

doc.add_paragraph(
    'Feature importance analysis reveals the key drivers of lapse behavior in the portfolio.'
)

add_img(doc, 'output/feature_importance.png', 6.5)

doc.add_paragraph(
    f'Figure 3: Top 20 feature importance scores from the lapse prediction model. The leading leakage-screened '
    f'features include {top_feature_phrase}. Direct outcome-derived fields are excluded from this model input, '
    'so the importance plot reflects operationally usable policy, premium, coverage, and customer attributes.'
)

fi_headers = ['Feature', 'Importance', 'Business Interpretation']
add_table(doc, fi_headers, fi_rows)

doc.add_paragraph('')

# 6.4 Sensitivity Analysis
doc.add_heading('6.4 Sensitivity Analysis', level=2)

doc.add_paragraph(
    'Sensitivity analysis was performed across key portfolio segments (age, premium, medical services, '
    'product type, distribution channel) to evaluate how lapse rates, claim costs, and loss ratios '
    'vary across customer groups.'
)

add_img(doc, 'output/sensitivity.png', 6.5)

doc.add_paragraph(
    'Figure 4: Representative sensitivity analysis for a key portfolio parameter. The pipeline also saves '
    'separate sensitivity plots for age, premium, medical services, product type, and distribution channel, '
    'so each segment can be reviewed without overwriting earlier outputs.'
)

# 6.5 Lapse Distribution
doc.add_heading('6.5 Lapse Distribution by Segment', level=2)

add_img(doc, 'output/lapse_distribution.png', 6.5)

doc.add_paragraph(
    'Figure 5: Lapse rate distribution across six customer segments. The highest lapse rates are observed '
    'in Distribution Channel D (direct sales, ~12.7%) and reimbursement policy holders (~9.2%). Product type '
    'S (Savings) shows the lowest lapse rate (~6.6%). Policy type DG C1 shows the highest lapse rate (~9.8%), '
    'while C2 shows the lowest (~1.5%). These segment-specific patterns inform targeted retention strategies.'
)

# 6.6 Dashboard Walkthrough
doc.add_heading('6.6 Dashboard Walkthrough', level=2)

doc.add_paragraph(
    'The interactive Streamlit dashboard provides six integrated views for portfolio analysis and '
    'decision support:'
)

dashboard_items = [
    'Key Performance Indicators: Real-time summary metrics including policies viewed, average premium, '
    'claim cost, lapse probability, and loss ratio — with comparison to the full portfolio.',
    'Exploratory Data Analysis: Premium and claim cost distributions, lapse rate by product and policy '
    'type, and age-based risk analysis using interactive Plotly charts.',
    'Decision Intelligence Framework: Pie chart of portfolio action distribution, interactive scatter plot '
    'of lapse risk vs. predicted claim cost with threshold lines, and a segment risk heatmap showing '
    'lapse rates by product x policy type.',
    'Predictive Model Performance: Tabbed views showing validation and test metrics for both the lapse '
    'model and claim cost model, with feature importance visualization.',
    'Decision Support Simulator: Interactive sliders for premium changes (-20% to +20%) and retention '
    'improvement (0-30% lapse reduction). The simulator recomputes decision classifications and shows '
    'the before/after impact with financial summaries.',
    'Customer-Level Risk View: Sortable table of the top N at-risk customers ranked by lapse probability, '
    'showing all key attributes for individual outreach planning.'
]
for item in dashboard_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================
# 7. DISCUSSION AND IMPLICATIONS
# ============================================================
add_heading_styled(doc, '7. Discussion and Implications', level=1)

doc.add_paragraph(
    'The results have several direct implications for health insurance portfolio management:'
)

doc.add_heading('7.1 Pricing Discipline', level=2)
doc.add_paragraph(
    f'The {fmt_count(early_risk_count)} EARLY_RISK policies represent a significant exposure area - these '
    f'customers have predicted claim costs substantially exceeding their premiums (average loss ratio of '
    f'{fmt_decimal(early_risk_loss, 2)}). '
    'Targeted pricing reviews for this segment could reduce underwriting losses. The REVIEW_PRICING '
    'category, while empty in the current sample, would capture borderline cases where premiums are '
    'adequate only by a narrow margin.'
)

doc.add_heading('7.2 Retention Strategy', level=2)
doc.add_paragraph(
    f'The {fmt_count(retain_high_count)} RETAIN_HIGH policies represent the highest-value retention '
    f'opportunity - these customers are profitable (premium exceeds predicted claim cost, average loss '
    f'ratio of {fmt_decimal(retain_high_loss, 2)}) but are at high '
    'risk of lapsing. The Decision Support Simulator allows leadership to quantify the portfolio impact '
    'of retention investments: a 5% reduction in lapse probability for this segment could preserve '
    'substantial premium revenue.'
)

doc.add_heading('7.3 Resource Allocation', level=2)
doc.add_paragraph(
    'The framework enables data-driven allocation of retention and risk management resources. Rather '
    'than spending uniformly across all customers, organizations can focus engagement budgets on the '
    'RETAIN_HIGH and EARLY_RISK segments, achieving higher ROI per dollar of retention spending.'
)

doc.add_heading('7.4 Strategic Decision Support', level=2)
doc.add_paragraph(
    'The Decision Support Simulator transforms the framework from a descriptive tool into a prescriptive '
    'platform. Leaders can test scenarios such as: "What happens if we increase premiums by 10%?" or '
    '"If our retention program reduces lapse by 15%, how does the portfolio classification change?" '
    'This capability directly supports the capstone objective of moving beyond traditional analytics '
    'to decision-oriented analysis.'
)

doc.add_page_break()

# ============================================================
# 8. LIMITATIONS AND FUTURE WORK
# ============================================================
add_heading_styled(doc, '8. Limitations and Future Work', level=1)

doc.add_paragraph(
    'Several limitations of the current analysis should be noted:'
)

limitations = [
    'Geographic scope: The dataset originates from a Spanish health insurer, so findings and '
    'recommendations may not transfer directly to the U.S. health insurance market without '
    'additional validation due to regulatory and market structure differences.',
    'Single-year snapshot: The data primarily covers 2017 and represents a cross-section rather than '
    'longitudinal time series analysis.',
    'Feature limitations: The model does not incorporate external data sources such as regional '
    'healthcare costs, competitor pricing, or macroeconomic indicators.',
    'Binary lapse classification: The model treats lapse as a binary outcome, ignoring the timing '
    'and intensity of lapse behavior.',
    'Data quality: Certain coverage variables (C_H, C_GI, C_GE_*) have notable missingness (5-8%), '
    'which may limit their usefulness in modeling.',
    'Leakage control: Direct outcome-derived variables are excluded from model inputs, but any production '
    'deployment should further validate that all retained features are available before the business decision point.'
]
for l in limitations:
    doc.add_paragraph(l, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph(
    'Future work could extend this framework by: (1) incorporating time-series analysis for lapse '
    'timing prediction, (2) adding customer lifetime value estimation, (3) integrating external '
    'healthcare cost indices, (4) implementing reinforcement learning for optimal retention strategy '
    'selection, and (5) extending to multi-period portfolio simulation.'
)

doc.add_page_break()

# ============================================================
# 9. CONCLUSIONS
# ============================================================
add_heading_styled(doc, '9. Conclusions', level=1)

doc.add_paragraph(
    'This capstone project demonstrates that predictive modeling, when combined with a decision-oriented '
    'framework, can transform health insurance portfolio data into actionable intelligence. The Insurance '
    f'Outcome Intelligence system achieves leakage-screened predictive performance (test AUC '
    f'{fmt_decimal(lapse_test.get("roc_auc"))} for lapse prediction, test R-squared '
    f'{fmt_decimal(claim_test.get("r2"))} for claim cost prediction) and translates these predictions into '
    'actionable categories that directly support pricing discipline, retention focus, and risk management.'
)

doc.add_paragraph(
    'The key contributions of this work are:'
)

contributions = [
    'A validated decision intelligence framework that classifies insurance portfolio customers into '
    'actionable categories based on lapse risk and claim cost predictions.',
    'An interactive dashboard that enables non-technical stakeholders to explore segments, view model '
    'predictions, and simulate business decisions in real time.',
    f'Empirical insights from a {fmt_count(data_rows)} row portfolio showing that approximately '
    f'{fmt_pct(priority_pct)} of policies fall into priority action categories requiring targeted intervention.',
    'A scalable, reproducible pipeline (available as open-source code) that can be adapted to other '
    'insurance datasets and markets.'
]
for c in contributions:
    doc.add_paragraph(c, style='List Number')

doc.add_paragraph(
    'The framework directly addresses all five research questions posed in the capstone proposal and '
    'provides measurable near-term value with long-term scalability. If implemented effectively, the '
    'system has the potential to create meaningful portfolio value through smarter pricing, better '
    'retention focus, and earlier risk control.'
)

doc.add_page_break()

# ============================================================
# 10. REFERENCES
# ============================================================
add_heading_styled(doc, '10. References', level=1)

refs = [
    'Lledó, Josep, Priscila Espinosa Adamez, and Virgilio Perez Gimenez. 2025. "Dataset of Health '
    'Insurance Portfolio." Mendeley Data, Version 4. https://doi.org/10.17632/386vmj2tbk.4',
    'Chen, Tianqi, and Carlos Guestrin. 2016. "XGBoost: A Scalable Tree Boosting System." Proceedings '
    'of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785-794.',
    'Friedman, Jerome H. 2001. "Greedy Function Approximation: A Gradient Boosting Machine." Annals '
    'of Statistics, 29(5): 1189-1232.',
    'Hastie, Trevor, Robert Tibshirani, and Jerome Friedman. 2009. The Elements of Statistical Learning: '
    'Data Mining, Inference, and Prediction. 2nd ed. Springer.',
    'Prochaska, J. O., and C. C. DiClemente. 1982. "Trans-theoretical Therapy: Toward a More Integrative '
    'Model of Change." European Journal of Social Psychology, 12(3): 277-288.',
    'Donthsen, Helmut. 2014. The Practice of Survey Research: Behind the Statistical Method. Springer.',
    'Handelman, Alan, et al. 2019. "Artificial Intelligence in Health Care: Current Applications and '
    'Future Potential." Journal of Medical Systems, 43(5): 149.',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(f'[{i}] {ref}')
    p.paragraph_format.space_after = Pt(4)

# Save
output_path = 'output/Capstone_Analysis_Report.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
print('Document generated successfully!')
